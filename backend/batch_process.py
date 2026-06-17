import os, sys, re, json, csv, time, pickle
from pathlib import Path
from datetime import datetime, timedelta
from math import ceil

# Ensure imports work from both Automatizacion-Ciclo/backend/ and root backend/
_BACKEND_DIR = str(Path(__file__).parent)
if _BACKEND_DIR not in sys.path:
    sys.path.insert(0, _BACKEND_DIR)

sys.stdout.reconfigure(line_buffering=True)
sys.stderr.reconfigure(line_buffering=True)

from services.phone_service import normalize_phone, extract_phone_from_text, deduplicate_phones
from services.gender_service import infer_gender, infer_gender_from_text
from services.experience_service import get_top_3_experiences
from services.rules_service import apply_rules

from app.services.llm_service import extract_fields
from app.services.cv_extractor import EXTRACTION_SCHEMA
import asyncio

CV_ROOT = Path(os.getenv("CV_ROOT", "/mnt/c/Users/ferna/OneDrive - Universidad de Chile/Documentos/Documentos Ciclo/Liceo Narciso Tondreau/concurso10434"))
MINEDUC_CSV = Path(os.getenv("MINEDUC_CSV", "/home/fernandosilvav/Proyecto-Prueba/backend/mineduc_matricula.csv"))
OUTPUT_DIR = Path(os.getenv("OUTPUT_DIR", "/home/fernandosilvav/Proyecto-Prueba/backend/outputs"))
STATE_FILE = Path(os.getenv("STATE_FILE", os.path.join(str(OUTPUT_DIR), ".batch_state.pkl")))

MAX_TEXT_PER_CANDIDATE = 12000

STATUS_FILE = OUTPUT_DIR / "batch_status.json"
BATCH_START_TIME = time.time()

OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

def write_status(current, total, candidate_id, name="?", status_msg="", flags=None):
    elapsed = time.time() - BATCH_START_TIME
    avg_per_candidate = elapsed / max(current, 1)
    remaining_est = avg_per_candidate * (total - current)
    eta_str = str(timedelta(seconds=int(remaining_est))) if current > 0 else "?"
    status = {
        "timestamp": datetime.now().isoformat(),
        "current": current,
        "total": total,
        "progress_pct": round(current / total * 100, 1) if total > 0 else 0,
        "candidate_id": candidate_id,
        "candidate_name": name,
        "status": status_msg,
        "elapsed_seconds": round(elapsed, 1),
        "elapsed_str": str(timedelta(seconds=int(elapsed))),
        "estimated_remaining": eta_str,
        "flags": flags or [],
    }
    with open(STATUS_FILE, "w") as f:
        json.dump(status, f, indent=2, ensure_ascii=False)

COLUMNS = [
    "NOMBRES", "APELLIDOS", "RUT", "GENERO",
    "TELEFONO_FIJO", "TELEFONO_CELULAR", "NACIONALIDAD",
    "TITULO_PROFESIONAL", "TITULO_ACADEMICO_1", "TITULO_ACADEMICO_2",
    "EXP1_EE", "EXP1_DURACION", "EXP1_MATRICULAS",
    "EXP2_EE", "EXP2_DURACION", "EXP2_MATRICULAS",
    "EXP3_EE", "EXP3_DURACION", "EXP3_MATRICULAS",
]

print("=" * 60)
print("BATCH PROCESSOR - concurso10434")
print(f"Candidates root: {CV_ROOT}")
print("=" * 60)

print("\n[1/6] Cargando MinEduc...")
me_lookup = {}
with open(MINEDUC_CSV, encoding="utf-8") as f:
    reader = csv.DictReader(f)
    for row in reader:
        name = row.get("NOM_RBD", "").strip()
        if name:
            normalized = re.sub(r"[^a-z0-9áéíóúñ\s]", "", name.lower())
            normalized = re.sub(r"\s+", " ", normalized).strip()
            mat_total = int(row.get("MAT_TOTAL", 0) or 0)
            existing = me_lookup.get(normalized)
            if existing is None or mat_total > int(existing.get("MAT_TOTAL", 0) or 0):
                me_lookup[normalized] = row

print(f"   {len(me_lookup)} establecimientos únicos cargados")

print("\n[2/6] Escaneando candidatos...")
candidate_dirs = sorted([d for d in CV_ROOT.iterdir() if d.is_dir()])
print(f"   {len(candidate_dirs)} candidatos encontrados")

def _extract_doc_text(filepath):
    import subprocess, html
    raw_bytes = filepath.read_bytes()

    text = ""
    try:
        decoded = raw_bytes.decode("utf-8", errors="replace")
        # NEW: Find real HTML content by searching for <html tag anywhere
        html_start = decoded.lower().find("<html")
        if html_start >= 0:
            html_end = decoded.lower().rfind("</html>")
            if html_end > html_start:
                real_html = decoded[html_start:html_end + 7]
            else:
                body_end = decoded.lower().rfind("</body>")
                real_html = decoded[html_start:body_end + 7] if body_end > html_start else decoded[html_start:html_start + 100000]
            import re as _re_html, html as _html_mod
            text = _re_html.sub(r"<[^>]+>", " ", real_html)
            text = _re_html.sub(r"\s+", " ", text).strip()
            text = _html_mod.unescape(text)
        else:
            # existing detection: check density and html tags in first 2000 chars
            sample = decoded[:2000]
            alpha_chars = sum(1 for c in sample if c.isalpha())
            density = alpha_chars / max(len(sample), 1)
            if density > 0.25 and ("<html" in decoded.lower() or "<p" in decoded.lower() or "<table" in decoded.lower() or "<div" in decoded.lower()):
                import re
                text = re.sub(r"<[^>]+>", " ", decoded)
                text = re.sub(r"\s+", " ", text).strip()
                text = html.unescape(text)
    except:
        pass

    if len(text) >= 50:
        return text

    try:
        result = subprocess.run(["antiword", str(filepath)], capture_output=True, text=True, timeout=10)
        if result.returncode == 0 and result.stdout.strip():
            text = result.stdout.strip()
            lines = [l for l in text.split("\n") if not l.startswith("convert ")]
            text = "\n".join(lines).strip()
            if len(text) >= 50:
                return text
    except FileNotFoundError:
        pass

    try:
        result = subprocess.run(["catdoc", str(filepath)], capture_output=True, text=True, timeout=10)
        if result.returncode == 0 and result.stdout.strip():
            return result.stdout.strip()
    except FileNotFoundError:
        pass

    try:
        import olefile
        if olefile.isOleFile(str(filepath)):
            ole = olefile.OleFileIO(str(filepath))
            if ole.exists('WordDocument'):
                stream = ole.openstream('WordDocument')
                data = stream.read()
                import re
                decoded_text = re.sub(rb'[^\x20-\x7E\x0A\x0D\xC0-\xFF\xE1\xE9\xED\xF3\xFA\xF1\xFC]', b' ', data)
                decoded_text = decoded_text.decode('latin-1', errors='replace')
                decoded_text = re.sub(r'\s+', ' ', decoded_text).strip()
                if len(decoded_text) >= 50:
                    ole.close()
                    return decoded_text
            ole.close()
    except ImportError:
        pass

    return ""


def extract_candidate_text(cand_dir):
    text_parts = []
    files = sorted(cand_dir.iterdir())
    for f in files:
        ext = f.suffix.lower()
        try:
            text = ""
            if ext == ".docx":
                from docx import Document
                doc = Document(str(f))
                _parts = [p.text for p in doc.paragraphs if p.text.strip()]
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            cell_text = cell.text.strip()
                            if cell_text:
                                _parts.append(cell_text)
                for section in doc.sections:
                    if section.header:
                        for p in section.header.paragraphs:
                            if p.text.strip():
                                _parts.append(p.text.strip())
                    if section.footer:
                        for p in section.footer.paragraphs:
                            if p.text.strip():
                                _parts.append(p.text.strip())
                text = "\n".join(_parts)
            elif ext == ".pdf":
                import pdfplumber
                with pdfplumber.open(str(f)) as pdf:
                    pages = []
                    for page in pdf.pages:
                        t = page.extract_text()
                        if t:
                            pages.append(t)
                    text = "\n".join(pages)
            elif ext == ".doc":
                text = _extract_doc_text(f)
            if text.strip():
                text_parts.append(f"--- {f.name} ---\n{text}")
        except Exception as e:
            pass
    combined = "\n\n".join(text_parts)
    SMART_QUOTES = {
        '\u201c': '"', '\u201d': '"', '\u201e': '"',
        '\u2018': "'", '\u2019': "'", '\u201a': "'",
        '\u00ab': '"', '\u00bb': '"',
    }
    for bad, good in SMART_QUOTES.items():
        combined = combined.replace(bad, good)
    return combined[:MAX_TEXT_PER_CANDIDATE]

BATCH_EXTRACTION_PROMPT = """Eres un extractor de datos de currículums chilenos. Analiza el CV y extrae la información solicitada.

Reglas importantes:
- Si un campo no aparece explícitamente en el CV, déjalo VACÍO (no pongas 'NO ENCONTRADO', solo déjalo como string vacío o null).
- NOMBRES y APELLIDOS por separado. Si el CV dice "Nombre: Juan Pérez García", NOMBRES="Juan", APELLIDOS="Pérez García"
- RUT con puntos y guión si aparece así
- Para teléfonos: extrae el número TAL CUAL aparece, sin modificar formato.
- NACIONALIDAD: casi siempre "Chilena"
- EXPERIENCIA LABORAL: extrae TODAS las experiencias que encuentres"""


def extract_with_llm(candidate_id, all_text, retries=3):
    try:
        result = asyncio.run(extract_fields(
            text=all_text,
            schema=EXTRACTION_SCHEMA,
            model="accurate",
            fallback_schema=EXTRACTION_SCHEMA,
            fallback_model="accurate",
            prompt_override=BATCH_EXTRACTION_PROMPT,
        ))
        return result if result else None
    except Exception as e:
        print(f"      Error en extracción: {e}")
        return None

def build_row(candidate_id, extracted):
    if not extracted:
        return None
    row = {col: "NO ENCONTRADO" for col in COLUMNS}
    row["RUT"] = candidate_id
    for src in ["NOMBRES", "APELLIDOS", "TELEFONO_FIJO", "TELEFONO_CELULAR", "NACIONALIDAD", "TITULO_PROFESIONAL", "TITULO_ACADEMICO_1", "TITULO_ACADEMICO_2"]:
        val = extracted.get(src)
        if val and str(val).strip() not in ("", "null", "None"):
            row[src] = str(val).strip()
    return row

print("\n[3/6] Extrayendo datos con Gemini...")

results = []
if STATE_FILE.exists():
    with open(STATE_FILE, "rb") as f:
        results = pickle.load(f)
    print(f"   Estado recuperado: {len(results)} candidatos ya procesados")

processed_ids = {r["id"] for r in results}
remaining = [d for d in candidate_dirs if d.name not in processed_ids]
sample_size_str = os.getenv("SAMPLE_SIZE", "")
if sample_size_str.isdigit():
    sample_size = int(sample_size_str)
    remaining = remaining[:sample_size]
    print(f"   Modo muestra: limitado a {sample_size} candidatos")
print(f"   Restantes: {len(remaining)} de {len(candidate_dirs)}")

for idx, cand_dir in enumerate(remaining):
    candidate_id = cand_dir.name
    current_num = len(results) + 1
    total = len(candidate_dirs)
    tick = time.time()

    write_status(current_num, total, candidate_id, status_msg="extrayendo texto...")
    print(f"\n   [{current_num}/{total}] RUT {candidate_id}", end=" ")
    sys.stdout.flush()

    text = extract_candidate_text(cand_dir)
    if not text.strip():
        elapsed = time.time() - tick
        print(f"SIN TEXTO ({elapsed:.1f}s)")
        row = {col: "NO ENCONTRADO" for col in COLUMNS}
        row["RUT"] = candidate_id
        flags = apply_rules(row, {}, COLUMNS)
        results.append({"id": candidate_id, "row": row, "extracted": None, "flags": flags})
        write_status(current_num, total, candidate_id, status_msg="sin texto")
        continue

    write_status(current_num, total, candidate_id, status_msg="enviando a Gemini...")
    extracted = extract_with_llm(candidate_id, text)
    if extracted is None:
        elapsed = time.time() - tick
        print(f"FALLÓ Gemini ({elapsed:.1f}s)")
        row = {col: "NO ENCONTRADO" for col in COLUMNS}
        row["RUT"] = candidate_id
        flags = apply_rules(row, {}, COLUMNS)
        results.append({"id": candidate_id, "row": row, "extracted": None, "flags": flags})
        write_status(current_num, total, candidate_id, status_msg="falló Gemini")
        continue

    row = build_row(candidate_id, extracted)
    if row is None:
        elapsed = time.time() - tick
        print(f"SIN DATOS ({elapsed:.1f}s)")
        row = {col: "NO ENCONTRADO" for col in COLUMNS}
        row["RUT"] = candidate_id
        flags = apply_rules(row, {}, COLUMNS)
        results.append({"id": candidate_id, "row": row, "extracted": None, "flags": flags})
        write_status(current_num, total, candidate_id, status_msg="sin datos extraídos")
        continue

    write_status(current_num, total, candidate_id, status_msg="post-procesando...")

    if row["GENERO"] == "NO ENCONTRADO":
        row["GENERO"] = infer_gender(row["NOMBRES"])
    if row["GENERO"] == "NO ENCONTRADO":
        row["GENERO"] = infer_gender_from_text(text)

    fijo_type, fijo_num = normalize_phone(row["TELEFONO_FIJO"])
    row["TELEFONO_FIJO"] = fijo_num if fijo_type != "NO ENCONTRADO" else "NO ENCONTRADO"
    cel_type, cel_num = normalize_phone(row["TELEFONO_CELULAR"])
    row["TELEFONO_CELULAR"] = cel_num if cel_type != "NO ENCONTRADO" else "NO ENCONTRADO"
    if row["TELEFONO_FIJO"] == "NO ENCONTRADO" and row["TELEFONO_CELULAR"] == "NO ENCONTRADO":
        phone_type, phone_num = extract_phone_from_text(text)
        if phone_type == "TELEFONO_CELULAR":
            row["TELEFONO_CELULAR"] = phone_num
        elif phone_type == "TELEFONO_FIJO":
            row["TELEFONO_FIJO"] = phone_num

    row = deduplicate_phones(row)

    exps = get_top_3_experiences(extracted.get("experiencia_laboral", []), me_lookup)
    for i, exp in enumerate(exps):
        row[f"EXP{i+1}_EE"] = exp.get("establecimiento", "NO ENCONTRADO")
        row[f"EXP{i+1}_DURACION"] = exp.get("duracion", "NO ENCONTRADO")
        row[f"EXP{i+1}_MATRICULAS"] = exp.get("matricula", "NO ENCONTRADO")
    for i in range(len(exps), 3):
        row[f"EXP{i+1}_EE"] = "NO ENCONTRADO"
        row[f"EXP{i+1}_DURACION"] = "NO ENCONTRADO"
        row[f"EXP{i+1}_MATRICULAS"] = "NO ENCONTRADO"

    # Fallback: if nombres still missing, try to extract from text
    if row["NOMBRES"] == "NO ENCONTRADO" and text.strip():
        import re as _re_email
        email_match = _re_email.search(r'([a-zA-Z][a-zA-Z._]+)@[a-zA-Z]', text)
        if email_match:
            local_part = email_match.group(1).lower()
            parts = _re_email.split(r'[._]', local_part)
            if parts and len(parts[0]) > 2:
                from services.gender_service import GENDER_MAP
                candidate_first = parts[0].title()
                if candidate_first.upper() in GENDER_MAP:
                    row["NOMBRES"] = candidate_first

    flags = apply_rules(row, extracted, COLUMNS)
    full_name = f"{row.get('NOMBRES','?')} {row.get('APELLIDOS','?')}".strip()
    fstr = ""
    if flags.get("orange"): fstr += " NARANJO"
    if flags.get("blue"): fstr += " AZUL"
    if flags.get("yellow"): fstr += " AMARILLO"
    elapsed = time.time() - tick
    print(f"{full_name}{fstr} ({elapsed:.1f}s)")

    results.append({"id": candidate_id, "row": row, "extracted": extracted, "flags": flags})
    write_status(current_num, total, candidate_id, full_name, status_msg="ok", flags=[k for k,v in flags.items() if v])

    if len(results) % 5 == 0:
        with open(STATE_FILE, "wb") as f:
            pickle.dump(results, f)
        print(f"   --- Checkpoint guardado ({len(results)} procesados) ---")
        time.sleep(0.5)

with open(STATE_FILE, "wb") as f:
    pickle.dump(results, f)

valid = [r for r in results if r["row"] is not None]
errors = [r for r in results if r["row"] is None]
print(f"\n[4/6] Procesados: {len(valid)} correctos, {len(errors)} errores")

print("\n[5/6] Generando Excel...")
from openpyxl import Workbook
from openpyxl.styles import PatternFill, Font, Alignment
from openpyxl.utils import get_column_letter

BLUE_FILL = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
YELLOW_FILL = PatternFill(start_color="FFD966", end_color="FFD966", fill_type="solid")
ORANGE_FILL = PatternFill(start_color="F4B183", end_color="F4B183", fill_type="solid")
GRAY_FILL = PatternFill(start_color="F2F2F2", end_color="F2F2F2", fill_type="solid")
HEADER_FILL = PatternFill(start_color="2F5496", end_color="2F5496", fill_type="solid")

wb = Workbook()
ws = wb.active
ws.title = "Resultados"

for ci, cn in enumerate(COLUMNS, 1):
    c = ws.cell(row=1, column=ci, value=cn)
    c.fill = HEADER_FILL
    c.font = Font(color="FFFFFF", bold=True)
    c.alignment = Alignment(horizontal="center")

for ri, r in enumerate(valid, 2):
    row_data = r["row"]
    flags = r["flags"]
    for ci, cn in enumerate(COLUMNS, 1):
        val = row_data.get(cn, "NO ENCONTRADO")
        cell = ws.cell(row=ri, column=ci, value=val)
        if flags.get("orange"):
            cell.fill = ORANGE_FILL
        elif flags.get("blue"):
            cell.fill = BLUE_FILL
            cell.font = Font(color="FFFFFF")
        elif flags.get("yellow"):
            cell.fill = YELLOW_FILL
        elif val == "NO ENCONTRADO":
            gray_flags = flags.get("gray", {})
            if gray_flags.get(cn, True):
                cell.fill = GRAY_FILL

widths = {"NOMBRES":20,"APELLIDOS":22,"RUT":15,"GENERO":12,"TELEFONO_FIJO":16,"TELEFONO_CELULAR":16,"NACIONALIDAD":14,"TITULO_PROFESIONAL":32,"TITULO_ACADEMICO_1":32,"TITULO_ACADEMICO_2":32,"EXP1_EE":35,"EXP1_DURACION":22,"EXP1_MATRICULAS":14,"EXP2_EE":35,"EXP2_DURACION":22,"EXP2_MATRICULAS":14,"EXP3_EE":35,"EXP3_DURACION":22,"EXP3_MATRICULAS":14}
for ci, cn in enumerate(COLUMNS, 1):
    ws.column_dimensions[get_column_letter(ci)].width = widths.get(cn, 15)

output = OUTPUT_DIR / "resultado_concurso10434.xlsx"
wb.save(str(output))
print(f"   Excel: {output}")
print(f"   Filas: {len(valid)}")
orange_count = sum(1 for r in valid if r['flags'].get('orange'))
blue_count = sum(1 for r in valid if r['flags'].get('blue'))
yellow_count = sum(1 for r in valid if r['flags'].get('yellow'))
print(f"   Naranjas (revisión manual): {orange_count}")
print(f"   Azules (actualidad>3): {blue_count}")
print(f"   Amarillas (datos no descifrados): {yellow_count}")
print("=" * 60)
